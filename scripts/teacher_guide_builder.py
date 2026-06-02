"""Teacher's Guide DOCX 生成器（9 大段长文）。

骨架严格对齐官方 Spring Days(L1) 与 Visiting Scotland(L4) 样本：
  1. Lesson Guide (Overview)         — Basic Info / Vocab&Phonics Goal / Key Objectives
  2. Pre-Reading Support              — Warm up / Phonics Focus 4 steps / Vocabulary Preview
  3. During Reading Strategies        — Picture Walk per page + Reading Routine + Rereading
  4. Post-Reading Practice (Book Activities)
  5. Post-Reading Practice (Worksheet)
  6. Reading Check                    — Words / Fluency / Comprehension
  7. Portfolio Creation Task          — 3 options (Creative / Oral / Critical Thinking)
  8. Independent Reading
  9. Lesson Close

每段标题用 Heading 2，正文 Poppins 11pt + Alibaba PuHuiTi 中文 fallback。
内容由 ai_extractor 抽取的 outline 信息 + 模板字符串组合而成；可在 web 端再编辑。
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from parser import BookOutline


FONT_EN = "Poppins"
FONT_CN = "Alibaba PuHuiTi 2.0 65 Medium"


def build_teacher_guide(outline: BookOutline, out_path: Path) -> Path:
    doc = Document()
    _set_default_font(doc)
    _set_a4_margins(doc)

    # 总标题
    h_title = doc.add_heading(level=1)
    run = h_title.add_run(f"Teacher's Guide: {outline.title}")
    _font(run, size_pt=22, bold=True)

    # === 1. Lesson Guide (Overview) ===
    _heading(doc, "Lesson Guide (Overview)")

    _heading(doc, "Basic Info:", level=3)
    _para(doc, f"Book Title: {outline.title};")
    _para(doc, f"Level: {_level_label(outline.level)};")
    _para(doc, f"Time: {outline.lesson_time or '60 mins'}")

    _heading(doc, "Vocabulary & Phonics Goal:", level=3)
    vocab_str = _format_vocab(outline)
    _para(doc, f"Vocabulary: {vocab_str}; Primary Phonics Rule: {outline.phonics or '—'}")

    _heading(doc, "Key Objectives:", level=3)
    _para(doc, _build_objectives(outline))

    # === 2. Pre-Reading Support ===
    _heading(doc, "Pre-Reading Support")

    _heading(doc, "Warm up & Purpose:", level=3)
    _para(doc, f"Teacher Says: \"Today we are going to read a book called {outline.title}. "
               f"Let's find out what the story is about and what we can learn from it.\"")
    _para(doc, "Teacher Asks: \"What do you already know about this topic? "
               "What words or ideas come to mind when you hear the title?\"")
    _para(doc, "Expected Response: Students share prior knowledge and predictions.")

    _heading(doc, "Phonics Focus (Interactive Multi-Step):", level=3)
    phonics = outline.phonics or "vowel patterns from the story"
    _para(doc, f"Rule Identification: {phonics}")
    _para(doc, "Step 1: Sound Discovery (I Do): Teacher models the sound and segments the word.")
    _para(doc, "Step 2: Word Family Extension (We Do): Teacher and students practice related words together.")
    _para(doc, "Step 3: Text Scavenger Hunt (You Do): Students find target sounds in the book.")
    _para(doc, "Step 4: Movement Challenge: Students act out or clap each sound segment.")

    _heading(doc, "Vocabulary Preview:", level=3)
    for word in _vocab_words(outline)[:6]:
        _heading(doc, f"{word}:", level=4)
        _para(doc, "Teacher Action: Show a clear gesture or visual that represents the word.")
        _para(doc, f"Teacher Says: \"This is {word}. Say {word} with me: {word}.\"")
        _para(doc, f"Expected Response: Students say \"{word}\" and mimic the gesture.")

    # === 3. During Reading Strategies ===
    _heading(doc, "During Reading Strategies")

    _heading(doc, "Step 1: Picture Walk (Visual Only):", level=3)
    for i, page in enumerate(outline.pages[1:8], start=2):
        _heading(doc, f"Page {i}", level=4)
        scene_hint = (page.scene or page.text or "").split(".")[0]
        _para(doc, f"Teacher Action: Point to key visual elements on this page.")
        _para(doc, f"Teacher Says: \"Look at Page {i}. {page.text}\"")
        _para(doc, f"Teacher Asks: \"What is happening in this picture?\"")
        _para(doc, f"Expected Response: Students describe what they see.")
        _para(doc, f"Teacher Confirms/Expands: \"{scene_hint or page.text}.\"")

    _heading(doc, "Step 2: Reading Routine (Every Page):", level=3)
    _para(doc, "Purpose: Practice decoding, fluency, and comprehension.")
    _para(doc, "1. Teacher reads the sentence aloud once.")
    _para(doc, "2. Students echo read.")
    _para(doc, "3. Students read together aloud.")
    _para(doc, "4. Students act out the action when possible.")

    _heading(doc, "Reading Comprehension Questions:", level=3)
    rr_qs = getattr(outline, "_rr_questions", []) or []
    for q in rr_qs[:-1]:  # 跳过最后的开放题（放到 wrap-up）
        page_part = f" (P{q.get('page')})" if q.get("page") else ""
        _para(doc, f"- {q.get('q', '')}{page_part}")

    _heading(doc, "Step 3: Rereading for Automaticity:", level=3)
    _para(doc, "Round 1: Read with expression, matching the emotion of each page.")
    _para(doc, "Round 2: Read with rhythm; tap a foot to keep a steady pace.")

    # === 4. Post-Reading Practice (Book Activities) ===
    _heading(doc, "Post-Reading Practice (Book Activities)")
    _para(doc, "Use the post-reading book pages to consolidate learning. "
               "Walk through each activity, model the first item, then have students complete the rest independently.")

    # === 5. Post-Reading Practice (Worksheet) ===
    _heading(doc, "Post-Reading Practice (Worksheet)")
    ws_qs = getattr(outline, "_worksheet_questions", []) or []
    for i, ws in enumerate(ws_qs[:6], 1):
        _heading(doc, f"Activity {i}: {ws.get('title', 'Activity')}", level=3)
        instr = ws.get("instruction", "")
        if instr:
            _para(doc, f"Goal: {instr}")
        _para(doc, "Teacher Script: Model the first item, then have students complete the activity. "
                   "Walk around and give targeted feedback.")
        if ws.get("extra"):
            _para(doc, f"Note: {ws['extra']}")

    # === 6. Reading Check ===
    _heading(doc, "Reading Check")
    _para(doc, "Purpose: Check word understanding, reading fluency, and reading comprehension.")
    _para(doc, "Step 1: Words Recognition — Teacher points to the words one by one. "
               "Student reads the words. Teacher ticks correct ones; circles ones to revisit.")
    _para(doc, "Step 2: Reading Fluency — Student reads the book independently. "
               "Teacher listens and marks words the student is struggling with.")
    _para(doc, "Step 3: Reading Comprehension and Expression — Teacher asks the questions. "
               "Student answers independently. Teacher ticks the questions answered correctly.")

    # === 7. Portfolio Creation Task ===
    _heading(doc, "Portfolio Creation Task")
    _para(doc, "Purpose: Create visible evidence of learning.")
    _heading(doc, "Option 1: Creative Arts:", level=3)
    _para(doc, _portfolio_creative(outline))
    _heading(doc, "Option 2: Oral Performance:", level=3)
    _para(doc, _portfolio_oral(outline))
    _heading(doc, "Option 3: Critical Thinking:", level=3)
    _para(doc, _portfolio_critical(outline))

    # === 8. Independent Reading ===
    _heading(doc, "Independent Reading (Optional)")
    _para(doc, "Student Task: Choose 2 books from the library on a related theme.")
    _para(doc, "Teacher Prompts: \"Look at the pictures.\" \"Try to read.\" \"Tell me one thing about the book.\"")

    # === 9. Lesson Close ===
    _heading(doc, "Lesson Close")
    _heading(doc, "Summarize:", level=3)
    _para(doc, _build_objectives(outline))
    _heading(doc, "Final Wrap-up:", level=3)
    _para(doc, f"Teacher Says: \"Today we read {outline.title} together. "
               f"We learned new vocabulary, practiced the phonics rule, and shared our ideas. "
               f"Great job, everyone — see you next time!\"")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# ---------- 内容生成器 ----------
def _vocab_words(outline: BookOutline) -> list[str]:
    if outline.vocabulary_simple:
        return outline.vocabulary_simple
    if outline.vocabulary_mastery:
        return outline.vocabulary_mastery + outline.vocabulary_exposure
    return []


def _format_vocab(outline: BookOutline) -> str:
    return ", ".join(_vocab_words(outline)) or "—"


def _build_objectives(outline: BookOutline) -> str:
    words = _vocab_words(outline)
    word_str = ", ".join(words[:4]) if words else "the target vocabulary"
    grammar = outline.grammar_focus or "the target sentence frames"
    phonics = outline.phonics or "the target phonics rule"
    return (
        f"Students will be able to identify and use the vocabulary words {word_str}; "
        f"recognize the phonics rule ({phonics}); use the grammar pattern {grammar}; "
        f"answer comprehension questions about the text; and express their own ideas "
        f"using the vocabulary and patterns from the book."
    )


def _portfolio_creative(outline: BookOutline) -> str:
    words = ", ".join(_vocab_words(outline)[:4]) or "the target vocabulary"
    return (
        f"Create a poster, collage, or mini-book inspired by {outline.title}. "
        f"Include drawings or pictures that show the vocabulary words ({words}). "
        f"Label each picture and write one sentence about your work."
    )


def _portfolio_oral(outline: BookOutline) -> str:
    return (
        f"Memorize 2–3 sentences from {outline.title} and perform them aloud with "
        f"expression and gestures. Add one sentence of your own opinion about the story."
    )


def _portfolio_critical(outline: BookOutline) -> str:
    return (
        f"Compare the main characters in {outline.title} with someone you know in real life. "
        f"Use a Venn diagram to show what is similar and what is different. "
        f"Explain your reasoning in 2–3 sentences."
    )


# ---------- DOCX 工具 ----------
def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT_EN
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), FONT_CN)


def _set_a4_margins(doc: Document) -> None:
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)


def _heading(doc, text: str, level: int = 2) -> None:
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    sizes = {1: 22, 2: 16, 3: 13, 4: 12}
    _font(run, size_pt=sizes.get(level, 12), bold=True)


def _para(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _font(run, size_pt=11)


def _font(run, *, size_pt: int, bold: bool = False) -> None:
    run.font.name = FONT_EN
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), FONT_CN)


def _level_label(level: str) -> str:
    s = (level or "").strip()
    if not s:
        return "1"
    if s.lower().startswith("smart"):
        return "Smart"
    digits = "".join(ch for ch in s if ch.isdigit())
    return digits or s
