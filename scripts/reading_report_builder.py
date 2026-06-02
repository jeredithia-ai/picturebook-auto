"""Reading Report DOCX 生成器 v2.1 — 严格 1 页 + 顶格 + 课堂参与度方框。

排版完全复刻 VIPKID 官方模板（Desktop/L0_Book92/L0_Book92 reading report.docx
+ Desktop/LX_BookXX reading report Sample.docx），并满足以下口径：

  • 严格 1 页 A4 portrait（A4 portrait 8.27×11.69 in，margins 1.0 cm）
  • 表格 cell 段落顶格对齐（vertical = TOP）
  • 字号 11pt + 行距 1.2（紧凑但不挤）
  • 阅读表达星级用 ★ (U+2605) + 橙色 #E97A24（emoji ⭐ 在 Word/WPS 上会渲染成 *）
  • 课堂参与度 emoji 紧贴 label，每组后跟一个空白方框 ☐ 供学生打勾

题量梯度（与 picture-book-workflow QA 规则一致）:
  L0/L1/L2/Smart = 4 题, ★ + ★★ + ★★ + ★★★
  L3/L4/L5/L6    = 5 题, ★ + ★★ + ★★ + ★★ + ★★★
末尾 ★★★ 题为生活化拓展题（不带 (P#)），其余必须带 (P#)。
"""
from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import config
from parser import BookOutline
from text_format import _to_us_spelling, format_word_answer, format_sentence_answer


FONT_EN = "Poppins"
FONT_CN = "阿里巴巴普惠体 2.0 55 Regular"
EMOJI_FONT = "Segoe UI Emoji"

# 表格列宽（严格按 sample，dxa 即 1/20 pt）
COL_WIDTHS_DXA = [1428, 1307, 1293, 1246, 1369, 1308, 1292, 1243]
TABLE_WIDTH_DXA = sum(COL_WIDTHS_DXA)  # 10486

# 行高（twips，1440 twips = 1 inch）。设为 atLeast，由内容撑开但有最小值
# 经实测：A4 portrait 可用高约 10.98"，下列值能让 L0-L6 都恰好落在一页内
HEADER_ROW_TWIPS = 480        # 灰底 section header
VOCAB_ROW_TWIPS = 700
DIFF_ROW_TWIPS = 1500         # 4 行 11pt 标签/值（标题/字数/CEFR/语法）
PHONICS_ROW_TWIPS = 600
FLUENCY_ROW_TWIPS = 1300      # 故事正文行（atLeast，长内容自然撑开）
QUESTIONS_ROW_TWIPS = 2200    # 4-5 题（atLeast，长题自然撑开）
ENGAGE_ROW_TWIPS = 900

HEADER_FILL = "F1F1F1"

# 字号 / 行距口径（pt）
TITLE_PT = 16
NAME_PT = 11
SECTION_PT = 13
BODY_PT = 11
LINE_SPACING = 1.2
PARA_AFTER_PT = 1   # 段后空隙（题目 / 难度 / 故事正文统一）

# 星级颜色（橙红，对照 sample）
STAR_COLOR = "E97A24"

# Logo 资源（可被 config.BRAND_DIR 下的同名文件覆盖）
LOGO_PATH = config.BRAND_DIR / "dino_reading_logo.png"


# ============================================================================
# 公开入口
# ============================================================================
def build_reading_report(outline: BookOutline, out_path: Path) -> Path:
    doc = Document()
    _set_a4_portrait(doc)
    _set_default_style(doc)

    # 移除 Document() 自带的初始空段
    body = doc.element.body
    for p in list(body.iter(qn("w:p"))):
        body.remove(p)

    _build_title_paragraph(doc, outline)
    _build_name_date_paragraph(doc)
    _build_main_table(doc, outline)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# ============================================================================
# 标题 / 姓名段
# ============================================================================
def _build_title_paragraph(doc, outline: BookOutline) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.0

    title_str = f"阅读报告 {_level_label(outline.level)} - {outline.title}"
    run = p.add_run(title_str)
    _bind_run(run, ascii_font=FONT_EN, east_asia=FONT_CN, size_pt=TITLE_PT, bold=True)

    if LOGO_PATH.exists():
        logo_run = p.add_run()
        logo_run.add_picture(str(LOGO_PATH), width=Cm(5.0))


def _build_name_date_paragraph(doc) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    pf.right_indent = Cm(0.4)

    run = p.add_run("姓名: __________  日期: _________ 年 _________ 月 _________ 日")
    _bind_run(run, ascii_font="Arial Black", east_asia=FONT_CN, size_pt=NAME_PT, bold=True)


# ============================================================================
# 主表（12 行 × 8 列）
# ============================================================================
def _build_main_table(doc, outline: BookOutline) -> None:
    table = doc.add_table(rows=12, cols=8)
    table.autofit = False

    _set_tbl_w(table, TABLE_WIDTH_DXA)
    _set_tbl_borders(table)
    _set_tbl_grid(table, COL_WIDTHS_DXA)

    # v1.8.3：L5-L6 把"自然拼读"行替换为"构词法"行
    phonics_label = "构词法" if _is_morphology_level(outline.level) else "自然拼读"
    sections = [
        ("阅读难度", DIFF_ROW_TWIPS, _fill_difficulty),
        ("词汇掌握", VOCAB_ROW_TWIPS, _fill_vocab),
        (phonics_label, PHONICS_ROW_TWIPS, _fill_phonics),
        ("阅读流利度", FLUENCY_ROW_TWIPS, _fill_fluency),
        ("阅读表达", QUESTIONS_ROW_TWIPS, _fill_questions),
        ("课堂参与度", ENGAGE_ROW_TWIPS, _fill_engagement),
    ]

    for sec_idx, (label, content_h, content_filler) in enumerate(sections):
        header_row = table.rows[sec_idx * 2]
        content_row = table.rows[sec_idx * 2 + 1]

        # ---- header row：合并 8 列 + 灰底 + 居中粗体 ----
        _set_row_height(header_row, HEADER_ROW_TWIPS)
        header_cell = _merge_row(header_row)
        _shade_cell(header_cell, HEADER_FILL)
        _vert_center(header_cell)
        _clear_and_fill_text(
            header_cell, label,
            size_pt=SECTION_PT, bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

        # ---- content row：高度 + 由 filler 决定是否合并 ----
        _set_row_height(content_row, content_h)
        content_filler(content_row, outline)


# ----------------------------------------------------------------------------
# Section fillers
# ----------------------------------------------------------------------------
def _fill_difficulty(row, outline: BookOutline) -> None:
    cell = _merge_row(row)
    _vert_top(cell)
    _clear_paragraphs(cell)

    # Reader Type 按 Level 强制映射（v1.8）
    # - L0 (Smart) / L1 / L2 是固定文本
    # - L3-L6 = "Fiction" / "Non-Fiction"（取 outline.fiction_type 或 outline.reader_type）
    reader_type = _default_reader_type(outline)

    # v2.0：按官方 L5-1 实测样本，词汇难度只显示短码（A2/B1），不带 "CEFR" 前缀和 Lexile
    vocab_code = _default_cefr_short_code(outline)

    pairs = [
        ("类型：", reader_type),
        ("阅读字数：", str(outline.total_words or "—")),
        ("词汇难度：", vocab_code),
        ("语法难度：", _normalize_grammar_cn(outline.grammar_focus) or "—"),
    ]
    for i, (label, value) in enumerate(pairs):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(PARA_AFTER_PT)
        p.paragraph_format.line_spacing = LINE_SPACING
        r_lbl = p.add_run(label)
        _bind_run(r_lbl, FONT_EN, FONT_CN, size_pt=BODY_PT, bold=True)
        r_val = p.add_run(value)
        _bind_run(r_val, FONT_EN, FONT_CN, size_pt=BODY_PT, bold=False)


def _fill_vocab(row, outline: BookOutline) -> None:
    """8 cells：4 词放 c0/c2/c4/c6，c1/c3/c5/c7 留空白打勾位。"""
    words = _vocab_words_for_rr(outline)
    for col_idx in range(8):
        cell = row.cells[col_idx]
        _vert_top(cell)
        _clear_paragraphs(cell)
        if col_idx % 2 == 0:
            word_idx = col_idx // 2
            text = words[word_idx] if word_idx < len(words) else ""
            if text:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(text)
                _bind_run(run, FONT_EN, FONT_CN, size_pt=BODY_PT, bold=False)


def _fill_phonics(row, outline: BookOutline) -> None:
    """自然拼读（L0-L4）/ 构词法（L5-L6）。"""
    cell = _merge_row(row)
    _vert_top(cell)
    _clear_paragraphs(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING

    if _is_morphology_level(outline.level):
        text = _normalize_morphology(outline.phonics, outline)
    else:
        text = _normalize_phonics(outline.phonics)
    run = p.add_run(text)
    _bind_run(run, FONT_EN, FONT_CN, size_pt=BODY_PT, bold=False)


def _fill_fluency(row, outline: BookOutline) -> None:
    cell = _merge_row(row)
    _vert_top(cell)
    _clear_paragraphs(cell)

    body_text = " ".join(
        page.text.strip()
        for page in outline.pages
        if page.page_type == "story" and page.text and page.text.strip()
    )
    # 故事内容超 250 字时再降到 10pt 进一步保证 1 页
    body_size = BODY_PT if len(body_text) <= 250 else 10

    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)
    p_title.paragraph_format.line_spacing = LINE_SPACING
    r_title = p_title.add_run(outline.title or "")
    _bind_run(r_title, FONT_EN, FONT_CN, size_pt=body_size, bold=False)

    if body_text:
        p_body = cell.add_paragraph()
        p_body.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_body.paragraph_format.space_before = Pt(0)
        p_body.paragraph_format.space_after = Pt(0)
        p_body.paragraph_format.line_spacing = LINE_SPACING
        r_body = p_body.add_run(body_text)
        _bind_run(r_body, FONT_EN, FONT_CN, size_pt=body_size, bold=False)


def _fill_questions(row, outline: BookOutline) -> None:
    cell = _merge_row(row)
    _vert_top(cell)
    _clear_paragraphs(cell)

    questions = _resolve_rr_questions(outline)

    for i, q in enumerate(questions, start=1):
        p = cell.paragraphs[0] if i == 1 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(PARA_AFTER_PT)
        p.paragraph_format.line_spacing = LINE_SPACING

        # 题干：美式拼写 + 独立 i 大写 + 首字母大写 + 句末问号
        import re as _re
        text = _to_us_spelling(q["q"].strip().rstrip(".！？!?")).strip()
        text = _re.sub(r"\bi\b", "I", text)
        text = _re.sub(r"\bi(['\u2019])", r"I\1", text)
        if text:
            text = text[0].upper() + text[1:]
        if text and text[-1] not in "?？.!":
            text += "?"
        page = q.get("page")
        stars = max(1, min(int(q.get("stars") or 1), 3))

        r_q = p.add_run(f"{i}. {text}")
        _bind_run(r_q, FONT_EN, FONT_CN, size_pt=BODY_PT, bold=False)

        if page is not None:
            r_p = p.add_run(f" (P{page})")
            _bind_run(r_p, FONT_EN, FONT_CN, size_pt=BODY_PT, bold=False)

        # 实心五角星 ★ (U+2605) — Poppins / Arial 都能正确渲染，避免 emoji ⭐ 在 Word/WPS
        # 渲染时 fallback 成 '*' 的问题。颜色用 sample 的橙红 #E97A24。
        r_s = p.add_run(" " + ("★" * stars))
        _bind_run(r_s, FONT_EN, FONT_CN, size_pt=BODY_PT, bold=False)
        _set_run_color(r_s, STAR_COLOR)


def _fill_engagement(row, outline: BookOutline) -> None:
    """课堂参与度：emoji + label + ☐ 复选框，3 组横向居中排列。"""
    cell = _merge_row(row)
    _vert_center(cell)
    _clear_paragraphs(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    items = [("😆", "Excellent"), ("😄", "Great"), ("🙂", "Good")]
    for idx, (emoji, label) in enumerate(items):
        r_e = p.add_run(emoji)
        _bind_run(r_e, EMOJI_FONT, FONT_CN, size_pt=BODY_PT + 2, bold=False)
        r_l = p.add_run(" " + label)
        _bind_run(r_l, FONT_EN, FONT_CN, size_pt=BODY_PT, bold=False)
        r_sp1 = p.add_run("  ")
        _bind_run(r_sp1, FONT_EN, FONT_CN, size_pt=BODY_PT, bold=False)
        r_box = p.add_run("☐")
        _bind_run(r_box, FONT_EN, FONT_CN, size_pt=BODY_PT + 4, bold=False)
        if idx < len(items) - 1:
            r_gap = p.add_run("          ")
            _bind_run(r_gap, FONT_EN, FONT_CN, size_pt=BODY_PT, bold=False)


# ============================================================================
# 数据规整
# ============================================================================
def _vocab_words_for_rr(outline: BookOutline) -> list[str]:
    """RR r4 严格 4 词：L0-2 用 mastery 第一行 4 词；L3-6 用 vocabulary 4 词。"""
    if outline.is_dual_vocab_level and outline.vocabulary_mastery:
        words = list(outline.vocabulary_mastery)
    elif outline.vocabulary_simple:
        words = list(outline.vocabulary_simple)
    elif outline.vocabulary_mastery:
        words = list(outline.vocabulary_mastery)
    else:
        words = []
    cleaned = []
    for w in words:
        ww = (w or "").strip().rstrip(",.;:").strip()
        if ww:
            # v1.8：词汇统一小写、美式拼写
            cleaned.append(_to_us_spelling(ww.lower()))
    while len(cleaned) < 4:
        cleaned.append("")
    return cleaned[:4]


def _resolve_rr_questions(outline: BookOutline) -> list[dict]:
    """读取 outline._rr_questions（AI 抽取 / 人工编辑后挂载），按口径星级补齐。

    口径（新规范）：
      • L0-L2: 4 题，3 颗星 ⭐⭐⭐ 不带 (P#)
      • L3:    5 题，3 颗星不带 (P#)；1-4 题带 (P#)
      • L4-L6: 5 题，**全部不带 (P#)**（新规范要求）

    若上游缺失，按 Level 题量给占位题。
    """
    raw = getattr(outline, "_rr_questions", None)
    dist = config.rr_question_distribution(outline.level)
    no_page = _no_page_numbers(outline.level)

    def _page_for(i: int, stars: int) -> int | None:
        if no_page:
            return None
        if stars == 3:
            return None
        return i + 2

    if raw and isinstance(raw, list) and len(raw) > 0:
        normalized: list[dict] = []
        for i, q in enumerate(raw[:len(dist)]):
            stars = dist[i]
            # 上游可显式给 page；若没给则默认 i+2（除非级别要求隐藏 or ⭐⭐⭐）
            if no_page or stars == 3:
                page = None
            else:
                page = q.get("page") or (i + 2)
            normalized.append({
                "q": str(q.get("q") or q.get("question") or "").strip(),
                "stars": stars,
                "page": page,
            })
        while len(normalized) < len(dist):
            i = len(normalized)
            stars = dist[i]
            normalized.append({
                "q": f"Question {i + 1}?",
                "stars": stars,
                "page": _page_for(i, stars),
            })
        return normalized

    return [
        {
            "q": f"Question {i + 1}?",
            "stars": stars,
            "page": _page_for(i, stars),
        }
        for i, stars in enumerate(dist)
    ]


def _no_page_numbers(level: str) -> bool:
    """L4-L6 报告中不标注绘本页码（新规范）。"""
    key = (str(level or "").strip().lower())
    if "smart" in key:
        return False
    digits = "".join(ch for ch in key if ch.isdigit())
    try:
        return int(digits) >= 4
    except ValueError:
        return False


def _default_reader_type(outline) -> str:
    """按 Level 强制映射 Reader Type（v1.8 规范）。

      Smart / L0 → Concept & Knowledge - Building Readers
      L1         → Patterned Narrative & Informational Readers
      L2         → Early Independent Genre-Exposure Readers
      L3 - L6    → Fiction  或  Non-Fiction   （取 outline.fiction_type / reader_type）

    用户在大纲中显式写了 reader_type 时，对 L0-L2 仍按级别覆盖；对 L3-L6 优先用 reader_type。
    """
    key = str(getattr(outline, "level", "") or "").strip().lower()
    if "smart" in key:
        return "Concept & Knowledge - Building Readers"
    digits = "".join(ch for ch in key if ch.isdigit())
    try:
        n = int(digits)
    except ValueError:
        return "Patterned Narrative & Informational Readers"
    if n == 0:
        return "Concept & Knowledge - Building Readers"
    if n == 1:
        return "Patterned Narrative & Informational Readers"
    if n == 2:
        return "Early Independent Genre-Exposure Readers"
    # L3 - L6: Fiction / Non-Fiction
    ft = (getattr(outline, "fiction_type", "") or getattr(outline, "reader_type", "") or "").strip()
    ft_low = ft.lower()
    if "non" in ft_low and "fic" in ft_low:
        return "Non-Fiction"
    if "fic" in ft_low:
        return "Fiction"
    return "Fiction"  # L3-L6 没填默认 Fiction


def _default_cefr_text(outline) -> str:
    """词汇难度长文本（带 CEFR 前缀）— 保留兼容旧调用。"""
    if getattr(outline, "cefr", "") and outline.cefr.strip():
        return f"CEFR {outline.cefr.strip()}"
    return f"CEFR {_default_cefr_short_code(outline)}"


def _default_cefr_short_code(outline) -> str:
    """词汇难度短码（按官方 L5-1 实测格式，无 CEFR 前缀，无 Lexile）：

      Smart → Pre-A1
      L0    → Pre-A1
      L1    → Pre-A1
      L2    → A1
      L3    → A1+
      L4    → A2
      L5    → A2     # 官方 L5-1 实测是 A2（不是 B1，官方更保守）
      L6    → B1

    若用户在大纲显式写了 cefr 字段，优先用大纲值。
    """
    if getattr(outline, "cefr", "") and outline.cefr.strip():
        # 移除可能的 "CEFR " 前缀
        cefr = outline.cefr.strip()
        if cefr.upper().startswith("CEFR "):
            cefr = cefr[5:].strip()
        return cefr
    key = str(getattr(outline, "level", "") or "").strip().lower()
    if "smart" in key:
        return "Pre-A1"
    digits = "".join(ch for ch in key if ch.isdigit())
    mapping = {
        0: "Pre-A1", 1: "Pre-A1", 2: "A1", 3: "A1+",
        4: "A2",     5: "A2",     6: "B1",
    }
    try:
        return mapping.get(int(digits), "A1")
    except ValueError:
        return "A1"


def _normalize_grammar_cn(raw: str) -> str:
    """语法难度统一显示中文时态名（v1.8.3 规则）。

    检测英文时态短语关键词 → 中文标准名：
      Simple present tense / present simple → 一般现在时态
      Simple past tense / past simple       → 一般过去时态
      Simple future tense                   → 一般将来时态
      Present continuous / present progressive → 现在进行时态
      Past continuous                       → 过去进行时态
      Present perfect                       → 现在完成时态
      Past perfect                          → 过去完成时态
      Modal verbs                            → 情态动词

    如果包含其他句型（"There is/are", "help + V" 等），按 "时态 + 句型" 拼接。
    用户已经填了中文则直接返回。
    """
    import re as _re
    s = (raw or "").strip()
    if not s:
        return ""
    # 已经含中文时态字眼 → 原样返回（去掉句末标点）
    if _re.search(r"[一二三过将完现][般去来成在]", s):
        return s.rstrip("。.；; ").strip()

    low = s.lower()
    tense_map = [
        ("present perfect", "现在完成时态"),
        ("past perfect", "过去完成时态"),
        ("present continuous", "现在进行时态"),
        ("present progressive", "现在进行时态"),
        ("past continuous", "过去进行时态"),
        ("past progressive", "过去进行时态"),
        ("simple future", "一般将来时态"),
        ("future simple", "一般将来时态"),
        ("simple past", "一般过去时态"),
        ("past simple", "一般过去时态"),
        ("past tense", "一般过去时态"),
        ("simple present", "一般现在时态"),
        ("present simple", "一般现在时态"),
        ("present tense", "一般现在时态"),
        ("modal verb", "情态动词"),
        ("imperative", "祈使句"),
    ]
    parts: list[str] = []
    found_tense = False
    for en, cn in tense_map:
        if en in low and cn not in parts:
            parts.append(cn)
            found_tense = True
            break
    # 句型附加
    if "there was" in low or "there were" in low:
        parts.append("there was/were 句型")
    elif "there is" in low or "there are" in low:
        parts.append("there is/are 句型")
    if "+ v" in low or "+ verb" in low:
        # 取动词搭配描述（如 "help + V" → "help + V 句型"）
        m = _re.search(r"(\w+)\s*\+\s*v", low)
        if m:
            parts.append(f"{m.group(1)} + V 句型")
    if not found_tense and not parts:
        # 没匹配任何时态 → 原文保留
        return s.rstrip("。.；; ").strip()
    return "；".join(parts)


def _is_morphology_level(level: str) -> bool:
    """L5 / L6 用构词法 (morphology)；其他级别用自然拼读 (phonics)。"""
    key = str(level or "").strip().lower()
    if "smart" in key:
        return False
    digits = "".join(ch for ch in key if ch.isdigit())
    try:
        n = int(digits)
    except ValueError:
        return False
    return n >= 5


# 常见后缀 / 前缀 → 含义（构词法兜底库，从故事词汇里自动检测）
_SUFFIX_MEANINGS = {
    "-ous":  "having/full of quality",
    "-ful":  "full of",
    "-less": "without",
    "-able": "can be / capable of",
    "-tion": "act/state of",
    "-ment": "result/action of",
    "-ness": "state/quality of",
    "-ly":   "in a ... manner (adverb)",
    "-er":   "person who / more",
    "-est":  "most (superlative)",
    "-ed":   "past tense / past participle",
    "-ing":  "continuous / gerund",
    "-y":    "having / characterized by",
    "-ish":  "somewhat / having quality of",
    "-en":   "to make / become",
    "-ity":  "quality / condition of",
}
_PREFIX_MEANINGS = {
    "un-":  "not / opposite of",
    "re-":  "again",
    "pre-": "before",
    "dis-": "not / opposite of",
    "in-":  "not",
    "im-":  "not",
    "non-": "not",
    "mis-": "wrongly / badly",
    "over-": "too much",
    "under-": "too little / beneath",
}


def _detect_morphology_in_words(words: list[str]) -> Optional[str]:
    """从词汇里检测最高频的构词法规则（后缀优先于前缀）。

    返回如 'suffix \"-ous\" (= having/full of quality, e.g. nervous, famous)' 的字符串。
    """
    if not words:
        return None
    words = [str(w or "").strip().lower() for w in words if str(w or "").strip()]

    # 后缀检测：词尾匹配
    for suf, meaning in _SUFFIX_MEANINGS.items():
        suf_letters = suf.lstrip("-")
        hits = [w for w in words if len(w) > len(suf_letters) + 1 and w.endswith(suf_letters)]
        if len(hits) >= 1:
            examples = ", ".join(hits[:3])
            return f'suffix "{suf}" (= {meaning}, e.g. {examples})'

    # 前缀检测：词头匹配
    for pre, meaning in _PREFIX_MEANINGS.items():
        pre_letters = pre.rstrip("-")
        hits = [w for w in words if len(w) > len(pre_letters) + 1 and w.startswith(pre_letters)
                # 排除 "income" 误判 "in-" 等：前缀后接元音/辅音规则太复杂，简化为只在词长>=4 时认
                and len(w) >= 4]
        if len(hits) >= 1:
            examples = ", ".join(hits[:3])
            return f'prefix "{pre}" (= {meaning}, e.g. {examples})'

    return None


def _normalize_morphology(raw: str, outline) -> str:
    """L5-L6 构词法格式化。

    - 用户填了 phonics 字段就用它（仍按 _normalize_phonics 规整 quote / 大小写）
    - 没填则从 outline 的 mastery / vocabulary 词里自动检测后缀/前缀
    - 都失败兜底为 'suffix "-ly" (= in a ... manner)'
    """
    if raw and str(raw).strip():
        return _normalize_phonics(raw)

    words: list[str] = []
    if getattr(outline, "vocabulary_mastery", None):
        words.extend(outline.vocabulary_mastery)
    if getattr(outline, "vocabulary_simple", None):
        words.extend(outline.vocabulary_simple)
    if getattr(outline, "vocabulary_exposure", None):
        words.extend(outline.vocabulary_exposure)

    auto = _detect_morphology_in_words(words)
    if auto:
        return auto
    return 'suffix "-ly" (= in a ... manner, adverb)'


def _normalize_phonics(raw: str) -> str:
    """规整 phonics 文本为 v1.8 sample 风格（不是句子，是词组）：

    规则：
      • 全小写（CEFR / PBL 等术语缩写保持大写）
      • 英文直双引号 \"...\"（自动把 curly quote 转回 straight quote）
      • 例词放在括号里：(friendship)
      • 不带句号
      • 多条规则用分号 + 空格分隔

    示例输入 → 输出：
      'Consonant blend FR (friendship).' → 'consonant blend \"fr\" (friendship)'
      'long o (snow ow)'                  → 'long \"o\" (snow ow)'
      'AI → /eɪ/: day, stay, play'        → 'ai → /eɪ/ (day, stay, play)'
      ''                                   → 'short vowel pattern'
    """
    import re

    if not raw:
        return "short vowel pattern"
    text = str(raw).strip().replace("\r\n", " ").replace("\n", " ")

    # 1. curly quote → straight quote
    text = (
        text.replace("\u201c", '"').replace("\u201d", '"')
            .replace("\u2018", '"').replace("\u2019", '"')
    )
    # 2. 去掉末尾句号
    text = text.rstrip(". ").strip()

    # 3. 整体小写化（但 CEFR / PBL 大写术语保留）— 直接全 lower，再把术语还原
    preserved = ["CEFR", "PBL"]
    text = text.lower()
    for term in preserved:
        text = re.sub(rf"\b{term.lower()}\b", term, text)

    # 4. ": word, word" → " (word, word)"
    m = re.search(r":\s*([a-zA-Z][a-zA-Z,\s]*[a-zA-Z])\s*$", text)
    if m:
        words = m.group(1).strip()
        text = text[:m.start()].rstrip() + f" ({words})"

    # 5. 若文本里完全没有双引号，尝试给"音素词"加双引号
    #    模式 a: <word(s)> <single token without quotes> ( <example words> )
    #    例如 "consonant blend fr (friendship)" → consonant blend "fr" (friendship)
    if '"' not in text:
        m2 = re.search(r"^(.*\S)\s+(\S+)\s*\(([^)]+)\)\s*$", text)
        if m2:
            head, token, examples = m2.group(1), m2.group(2), m2.group(3)
            # 防止把已经是 "blend"/"vowel" 等普通修饰词加引号 — 用启发式：
            # token 仅 1-4 个字母且不在常见 stopwords 里
            stop = {"blend", "vowel", "consonant", "long", "short", "digraph",
                    "diphthong", "the", "and"}
            if token.lower() not in stop and 1 <= len(token) <= 5 and token.replace("+", "").isalpha():
                text = f'{head} "{token}" ({examples})'
    return text


def _level_label(level: str) -> str:
    """Level 0 / Smart / 1-6 → 标题里的 'Level X' 或 'Smart'."""
    s = str(level or "").strip()
    low = s.lower()
    if low.startswith("smart"):
        return "Smart"
    digits = "".join(ch for ch in s if ch.isdigit())
    return f"Level {digits or '1'}"


# ============================================================================
# python-docx XML 工具
# ============================================================================
def _bind_run(run, ascii_font: str, east_asia: str, size_pt: int, bold: bool) -> None:
    """每个 run 必须双绑 ascii + eastAsia 字体，强制规整。"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), east_asia)
    rFonts.set(qn("w:cs"), ascii_font)

    half_pt = str(int(size_pt * 2))
    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rPr.append(sz)
    sz.set(qn("w:val"), half_pt)
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), half_pt)

    if bold:
        if rPr.find(qn("w:b")) is None:
            rPr.append(OxmlElement("w:b"))
        if rPr.find(qn("w:bCs")) is None:
            rPr.append(OxmlElement("w:bCs"))


def _set_default_style(doc) -> None:
    style = doc.styles["Normal"]
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    rFonts.set(qn("w:cs"), FONT_EN)


def _set_a4_portrait(doc) -> None:
    """A4 portrait 8.27×11.69 in，margins 1.0 cm（紧凑版式确保 1 页）"""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(0.9)
        section.bottom_margin = Cm(0.9)


def _set_tbl_w(table, width_dxa: int) -> None:
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(width_dxa))
    tblW.set(qn("w:type"), "dxa")


def _set_tbl_borders(table) -> None:
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for kind in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{kind}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "BFBFBF")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_tbl_grid(table, widths_dxa: list[int]) -> None:
    tbl = table._element
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is not None:
            tblPr.addnext(grid)
        else:
            tbl.insert(0, grid)
    for gc in list(grid.findall(qn("w:gridCol"))):
        grid.remove(gc)
    for w in widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)


def _set_row_height(row, twips: int) -> None:
    trPr = row._tr.get_or_add_trPr()
    h = trPr.find(qn("w:trHeight"))
    if h is None:
        h = OxmlElement("w:trHeight")
        trPr.append(h)
    h.set(qn("w:val"), str(twips))
    h.set(qn("w:hRule"), "atLeast")


def _merge_row(row):
    cells = row.cells
    merged = cells[0]
    for c in cells[1:]:
        merged = merged.merge(c)
    return merged


def _shade_cell(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)


def _vert_center(cell) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _vert_top(cell) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def _set_run_color(run, hex_color: str) -> None:
    rPr = run._element.get_or_add_rPr()
    color = rPr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rPr.append(color)
    color.set(qn("w:val"), hex_color)


def _clear_paragraphs(cell) -> None:
    paras = cell.paragraphs
    if not paras:
        return
    first = paras[0]
    for run in list(first.runs):
        first._element.remove(run._element)
    for p in paras[1:]:
        p._element.getparent().remove(p._element)


def _clear_and_fill_text(cell, text: str, *, size_pt: int = 14,
                         bold: bool = False,
                         align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    _clear_paragraphs(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    _bind_run(run, FONT_EN, FONT_CN, size_pt=size_pt, bold=bold)


# ============================================================================
# 兼容入口
# ============================================================================
def attach_rr_questions(outline: BookOutline, questions: list[dict]) -> None:
    setattr(outline, "_rr_questions", questions)
